"""
scenario_test.py — Minimal validation of deepagents for CodeMind capabilities.

Validates:
  1. Long-context: Word document parsing → entity extraction → DB storage
  2. Multi-turn conversation: memory persistence & skill following across turns
  3. Skill adherence: agent respects SKILL.md rules across turns

Shared resources (from CodeMind):
  - .env          — LLM endpoints & keys
  - docs/         — architecture docs & requirement .docx files
  - utils/paths.py — unified path resolution

Usage:
  pip install -r requirements.txt
  python scenario_test.py              # uses first .docx from docs/req/
  python scenario_test.py --docx path  # use specific document
  python scenario_test.py --synthetic  # generate test doc (no shared docs needed)
"""

from __future__ import annotations

import argparse
import asyncio
import json
import os
import sys
import textwrap
from pathlib import Path

PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT))

# Fix Unicode output on Windows terminals (GBK can't handle emoji/Chinese properly)
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")


# ═══════════════════════════════════════════════════════════════════════════════
# 1. Environment & paths
# ═══════════════════════════════════════════════════════════════════════════════

def load_dotenv() -> None:
    """Load .env (shared with CodeMind)."""
    from dotenv import load_dotenv as _load

    env_path = PROJECT_ROOT / ".env"
    if not env_path.exists():
        print("[WARN] .env not found. Set LLM_API_KEY, LLM_BASE_URL, LLM_MODEL_ID manually.")
        return

    _load(env_path)
    masked = os.getenv("LLM_API_KEY", "")
    display = masked[:8] + "..." if len(masked) > 8 else "(empty)"
    print(f"[OK] Loaded .env  (model={os.getenv('LLM_MODEL_ID', '?')}, key={display})")


def find_docx() -> str | None:
    """Find the first usable .docx from shared docs/req/."""
    from utils.paths import data_paths

    req_dir = data_paths.shared_req_dir()
    if not req_dir.exists():
        return None

    for f in sorted(req_dir.iterdir()):
        if f.suffix.lower() == ".docx" and not f.name.startswith("~"):
            return str(f)
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# 2. Synthetic test document (fallback when no shared docs available)
# ═══════════════════════════════════════════════════════════════════════════════

def create_test_docx() -> str:
    """Generate a minimal test requirement document. Used only as fallback."""
    from docx import Document

    doc = Document()

    for i, (name, size) in enumerate([("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)], 1):
        try:
            style = doc.styles[name]
            style.font.size = doc.shared.Pt(size)
            style.font.bold = True
        except Exception:
            pass

    doc.add_heading("XX电商平台用户需求规格说明书", level=1)
    doc.add_paragraph(
        "本文档描述了XX电商平台V3.0版本的核心用户需求，涵盖用户管理、订单系统、"
        "商品管理三大核心模块。系统建设目标：打造高性能、高可用的电商平台，支持日均百万级订单处理。"
    )

    doc.add_heading("用户管理模块", level=2)
    doc.add_paragraph("用户管理模块负责终端用户的注册、登录、权限控制、个人信息管理等功能。")

    doc.add_heading("用户注册与登录", level=3)
    doc.add_paragraph(
        "系统应支持手机号+验证码注册方式。性能要求：注册接口响应时间小于500ms，支持1000并发。"
    )

    doc.add_heading("角色与权限管理", level=3)
    doc.add_paragraph(
        "系统应支持基于RBAC的角色权限控制，权限粒度需达到按钮级别。"
        "默认提供管理员、运营、普通用户三种角色。"
    )

    doc.add_heading("订单系统模块", level=2)
    doc.add_paragraph("订单系统负责用户下单、支付、物流追踪、退换货等全生命周期管理。")

    doc.add_heading("购物车与订单创建", level=3)
    doc.add_paragraph(
        "下单接口响应时间小于200ms，支持5000并发。超卖场景需通过分布式锁保障库存一致性。"
    )

    doc.add_heading("订单支付", level=3)
    doc.add_paragraph("支持微信支付、支付宝、银行卡三种支付方式。支付超时时间15分钟。")

    doc.add_heading("商品管理模块", level=2)
    doc.add_paragraph("商品管理模块包括商品发布、SKU管理、库存管理、商品搜索等功能。")

    doc.add_heading("SKU与库存管理", level=3)
    doc.add_paragraph(
        "每个商品可配置多个SKU，每个SKU独立管理库存和价格。"
        "库存变更需记录日志，支持低库存预警（阈值默认10件）。"
    )

    out_path = PROJECT_ROOT / "data" / "test_req.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] Generated test document: {out_path}  ({out_path.stat().st_size} bytes)")
    return str(out_path)


# ═══════════════════════════════════════════════════════════════════════════════
# 3. Model & agent
# ═══════════════════════════════════════════════════════════════════════════════

def build_model():
    """Build ChatOpenAI model using same config as CodeMind's .env."""
    from langchain_openai import ChatOpenAI

    model_id = os.getenv("LLM_MODEL_ID", "deepseek-v4-flash")
    api_key = os.getenv("LLM_API_KEY", "")
    base_url = os.getenv("LLM_BASE_URL", "https://api.deepseek.com")
    timeout = int(os.getenv("LLM_TIMEOUT", "180"))

    if not api_key:
        print("[WARN] LLM_API_KEY not set — agent will fail at runtime")

    extra_body = {}
    if "deepseek" in model_id.lower():
        extra_body = {"thinking": {"type": "disabled"}}

    return ChatOpenAI(
        model=model_id,
        api_key=api_key,
        base_url=base_url,
        timeout=timeout,
        max_retries=1,
        temperature=0.0,
        max_tokens=8192,
        extra_body=extra_body,
    )


def build_agent():
    """Create the deep agent with tools, skills, memory, and checkpointer."""
    from deepagents import create_deep_agent
    from deepagents.backends import FilesystemBackend
    from langgraph.checkpoint.memory import MemorySaver

    from tools.entity_store import entity_store, init_db
    from tools.read_docx import read_docx
    from tools.sql_query import create_sql_query_tool

    init_db()
    print("[OK] Database initialized (entity_store)")

    model = build_model()
    checkpointer = MemorySaver()

    # Text2SQL tool — shares the same model instance
    query_reqmgmt = create_sql_query_tool()

    skills_dir = str(PROJECT_ROOT / "skills" / "req-parse")
    memory_file = str(PROJECT_ROOT / "memory" / "AGENTS.md")

    agent = create_deep_agent(
        model=model,
        tools=[read_docx, entity_store, query_reqmgmt],
        system_prompt=(
            "You are a requirements management assistant. "
            "You help users parse requirement documents (.docx), extract structured entities "
            "(products, requirement_models, requirement_items), store them to a database, and query them.\n\n"
            "## Entity Extraction Rules\n"
            "When parsing a requirement document, follow these mapping rules:\n\n"
            "**products (产品, _type='products')**:\n"
            "- The document title (first H1 / `#` heading) maps to a product\n"
            "- id format: `PROD-001` (auto-increment), name = the title text, description = overview paragraphs\n\n"
            "**requirement_models (需求模型 / RM, _type='requirement_models')**:\n"
            "- Each H2 / `##` heading section maps to a requirement_model\n"
            "- id format: `RM-001`, `RM-002`, etc.\n"
            "- name = `{id} {heading text}`, type = `user_requirement`, product_id = the product id\n"
            "- description = first body paragraph under the H2 heading\n\n"
            "**requirement_items (用户需求项 / IR, _type='requirement_items')**:\n"
            "- Each H3 / `###` heading subsection maps to a requirement_item\n"
            "- id format: `IR-001`, `IR-002`, etc. (sequential across the whole document)\n"
            "- name = same as id, title = the H3 heading text\n"
            "- priority = `中`, status = `未实现`\n"
            "- rm_id = the id of the parent requirement_model (the H2 that contains this H3)\n\n"
            "**Steps**: 1) read_docx → 2) extract entities following the rules above → "
            "3) entity_store operation='store' with the entities as JSON (each MUST have _type field) → "
            "4) entity_store operation='stats' to verify."
        ),
        skills=[skills_dir],
        memory=[memory_file],
        checkpointer=checkpointer,
        backend=FilesystemBackend(root_dir=str(PROJECT_ROOT), virtual_mode=True),
    )

    print(f"[OK] Agent created (model={model.model_name}, tools: read_docx, entity_store, query_reqmgmt)")
    return agent, checkpointer


# ═══════════════════════════════════════════════════════════════════════════════
# 4. Test runner
# ═══════════════════════════════════════════════════════════════════════════════

async def run_turn(agent, thread_id: str, user_message: str, turn_label: str, max_steps: int = 30) -> dict:
    """Run a single conversation turn with async streaming."""
    import time as _time

    config = {"configurable": {"thread_id": thread_id}}

    print(f"\n{'='*60}")
    print(f"  {turn_label}")
    print(f"  User: {user_message[:120]}{'...' if len(user_message) > 120 else ''}")
    print(f"{'='*60}")

    all_messages: list = []
    tool_calls_found: list[str] = []
    response_text = ""
    step = 0
    turn_start = _time.time()

    # Suppress [LOG] spam from InvocationLoggingHandler (prints via stdout, not logging)
    _real_stdout = sys.stdout
    sys.stdout = type("_LogFilter", (object,), {
        "write": lambda self, s: _real_stdout.write(s) if "[LOG]" not in str(s) else None,
        "flush": lambda self: _real_stdout.flush(),
        "__getattr__": lambda self, k: getattr(_real_stdout, k),
    })()

    print(f"\n  [Streaming — max {max_steps} steps]\n")

    input_data = {"messages": [{"role": "user", "content": user_message}]}

    try:
        async for chunk in agent.astream(input_data, config=config, stream_mode="updates"):
            step += 1
            if step > max_steps:
                print(f"  [HALT] Max steps ({max_steps}) reached")
                break

            ts = _time.strftime("%H:%M:%S")
            elapsed = _time.time() - turn_start

            for node_name, node_output in chunk.items():
                if node_output is None:
                    continue
                if "messages" not in node_output:
                    continue

                msgs = node_output["messages"]
                if hasattr(msgs, "value"):
                    msgs = msgs.value
                if not isinstance(msgs, (list, tuple)):
                    msgs = [msgs]

                show_header = False
                for msg in msgs:
                    if msg not in all_messages:
                        if not show_header:
                            print(f"  [{ts}] Step {step} (+{elapsed:.1f}s) | Node: {node_name}")
                            show_header = True
                        all_messages.append(msg)
                        msg_type = getattr(msg, "type", "?")

                        if hasattr(msg, "tool_calls") and msg.tool_calls:
                            for tc in msg.tool_calls:
                                name = tc.get("name", "unknown")
                                args = tc.get("args", {})
                                tool_calls_found.append(name)
                                args_str = json.dumps(args, ensure_ascii=False, indent=2)
                                if len(args_str) > 500:
                                    args_str = args_str[:500] + "\n    ..."
                                print(f"      [TOOL CALL] {name}")
                                print(f"      Args:\n{textwrap.indent(args_str, '        ')}")

                        elif msg_type == "ai":
                            content = str(getattr(msg, "content", ""))
                            if content:
                                if len(content) > 400:
                                    content = content[:400] + "..."
                                print(f"      [AI] {content}")

                        elif msg_type == "tool":
                            content = str(getattr(msg, "content", ""))
                            tool_name = getattr(msg, "name", "?")
                            if len(content) > 400:
                                content = content[:400] + "..."
                            print(f"      [TOOL RESULT] {tool_name}: {content}")

                print()
    except Exception as e:
        sys.stdout = _real_stdout
        print(f"\n  [ERROR] Turn failed after {_time.time() - turn_start:.1f}s: {e}")
        import traceback
        traceback.print_exc()
        return {"response": f"ERROR: {e}", "tools": tool_calls_found}

    sys.stdout = _real_stdout
    elapsed_total = _time.time() - turn_start
    print(f"  [DONE] Turn completed in {elapsed_total:.1f}s ({step} steps)")

    for msg in reversed(all_messages):
        if getattr(msg, "type", "") == "ai":
            content = str(getattr(msg, "content", ""))
            if content:
                response_text = content
                break

    print(f"  Tools called: {tool_calls_found if tool_calls_found else '(none)'}")
    print(f"\n  Final response preview:\n{textwrap.indent(response_text[:800], '    ')}")
    if len(response_text) > 800:
        print(f"    ... (truncated, total {len(response_text)} chars)")

    return {"response": response_text, "tools": tool_calls_found}


# ═══════════════════════════════════════════════════════════════════════════════
# 5. Verification
# ═══════════════════════════════════════════════════════════════════════════════

def verify_skill_adherence(tool_calls_by_turn: dict[str, list[str]]) -> dict:
    """Check that the agent used the right tools."""
    checks = {}
    for turn_label, tools in tool_calls_by_turn.items():
        checks[f"{turn_label}_used_read_docx"] = "read_docx" in tools
        checks[f"{turn_label}_used_entity_store"] = "entity_store" in tools
        checks[f"{turn_label}_used_query_reqmgmt"] = "query_reqmgmt" in tools
    return checks


def verify_database() -> dict:
    """Direct DB check — ground truth for entity storage."""
    import sqlite3

    from utils.paths import data_paths

    db_path = data_paths.reqmgmt_db()
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row

    products = [dict(r) for r in conn.execute("SELECT * FROM products").fetchall()]
    models = [dict(r) for r in conn.execute("SELECT * FROM requirement_models").fetchall()]
    items = [dict(r) for r in conn.execute("SELECT * FROM requirement_items").fetchall()]
    conn.close()

    return {
        "product_count": len(products),
        "model_count": len(models),
        "item_count": len(items),
        "products": products,
        "models": models,
        "items": items,
    }


def verify_response_indicators(responses: dict[str, str]) -> dict:
    """Check for key indicators in agent responses."""
    indicators = {}
    for turn_label, text in responses.items():
        t = text.lower()
        indicators[f"{turn_label}_mentions_product"] = "产品" in text or "product" in t
        indicators[f"{turn_label}_mentions_models"] = "需求模型" in text or "requirement_model" in t
        indicators[f"{turn_label}_mentions_items"] = "需求项" in text or "requirement_item" in t
        indicators[f"{turn_label}_mentions_storage"] = "入库" in text or "stored" in t or "inserted" in t
        indicators[f"{turn_label}_mentions_counts"] = "统计" in text or "count" in t or "total" in t
    return indicators


# ═══════════════════════════════════════════════════════════════════════════════
# 6. Main
# ═══════════════════════════════════════════════════════════════════════════════

async def main():
    parser = argparse.ArgumentParser(description="DeepMind — deepagents validation")
    parser.add_argument("--docx", help="Path to a .docx file to parse")
    parser.add_argument("--synthetic", action="store_true",
                        help="Generate a synthetic test document instead of using shared docs")
    args = parser.parse_args()

    print("=" * 60)
    print("  DeepMind — deepagents validation for CodeMind capabilities")
    print("=" * 60)

    # ── Env ──
    load_dotenv()

    # ── Document selection ──
    if args.docx:
        docx_path = args.docx
        if not Path(docx_path).exists():
            print(f"[FAIL] Document not found: {docx_path}")
            sys.exit(1)
        print(f"[OK] Using specified document: {docx_path}")
    elif args.synthetic:
        docx_path = create_test_docx()
    else:
        docx_path = find_docx()
        if docx_path:
            print(f"[OK] Using shared document: {docx_path}")
        else:
            print("[INFO] No shared docs found in docs/req/, generating synthetic test doc.")
            print("       Place .docx files in docs/req/ to use real documents.")
            docx_path = create_test_docx()

    # ── Agent ──
    agent, checkpointer = build_agent()

    import time
    thread_id = f"test-{int(time.time())}"

    # ── Turn 1: Parse & Store (tests skill following + long context) ──
    turn1_msg = (
        f"请解析需求文档 `{docx_path}`，严格按照 req-parse skill 的规则提取其中的产品(product)、"
        f"需求模型(requirement_model)和用户需求项(requirement_item)，然后存储到数据库。"
        f"完成后用 entity_store stats 确认入库结果。"
    )
    turn1 = await run_turn(agent, thread_id, turn1_msg, "Turn 1: Parse & Store (Skill + Long Context)")

    # ── Turn 2: Query & Filter (tests multi-turn memory) ──
    turn2_msg = (
        "刚才一共入库了多少条需求项(requirement_item)？"
        "按需求模型分类展示每个模型下的需求项数量。"
        "另外，把包含'性能'关键词的需求项筛选出来。"
    )
    turn2 = await run_turn(agent, thread_id, turn2_msg, "Turn 2: Query & Filter (Multi-turn Memory)")

    # ── Turn 3: Contextual follow-up (tests sustained cross-turn reference) ──
    turn3_msg = (
        "回到我们在第一轮解析的那个文档，其中'订单系统模块'下面有几个需求项？"
        "列出它们的具体内容。"
    )
    turn3 = await run_turn(agent, thread_id, turn3_msg, "Turn 3: Contextual Follow-up (Cross-turn Reference)")

    # ── Turn 4: Skill re-application (tests skill adherence after context buildup) ──
    turn4_msg = (
        "如果我修改了文档中的某个需求项（比如把'支付超时时间15分钟'改成'30分钟'），"
        "你会如何更新数据库？请说明步骤，不必实际执行。"
    )
    turn4 = await run_turn(agent, thread_id, turn4_msg, "Turn 4: Skill Reasoning (Adherence After Context)")

    # ── Turn 5: Natural language query (tests Text2SQL via query_reqmgmt tool) ──
    turn5_msg = (
        "使用 query_reqmgmt 工具查询：数据库里一共有多少个产品(product)？"
        "列出所有需求模型(requirement_model)的名称和类型。"
    )
    turn5 = await run_turn(agent, thread_id, turn5_msg, "Turn 5: Text2SQL Basic Query")

    # ── Turn 6: Complex Text2SQL query (tests SQL generation from NL) ──
    turn6_msg = (
        "使用 query_reqmgmt 工具查询："
        "每个需求模型下各有多少需求项(requirement_item)？按数量从多到少排列。"
        "另外，查询优先级为'高'的需求项有哪些？"
    )
    turn6 = await run_turn(agent, thread_id, turn6_msg, "Turn 6: Text2SQL Aggregate & Filter")

    # ── Verification ──
    print(f"\n{'='*60}")
    print("  Verification Summary")
    print(f"{'='*60}")

    responses = {
        "turn1": turn1["response"],
        "turn2": turn2["response"],
        "turn3": turn3["response"],
        "turn4": turn4["response"],
        "turn5": turn5["response"],
        "turn6": turn6["response"],
    }
    tool_calls_by_turn = {
        "turn1": turn1["tools"],
        "turn2": turn2["tools"],
        "turn3": turn3["tools"],
        "turn4": turn4["tools"],
        "turn5": turn5["tools"],
        "turn6": turn6["tools"],
    }

    # DB ground truth
    db_state = verify_database()
    print("\n  [Database Ground Truth]")
    print(f"    products:           {db_state['product_count']}")
    print(f"    requirement_models: {db_state['model_count']}")
    print(f"    requirement_items:  {db_state['item_count']}")
    for m in db_state["models"]:
        print(f"      {m['name']}")
    for it in db_state["items"]:
        print(f"      {it['name']}: {it['title']}")

    # Skill adherence
    skill_checks = verify_skill_adherence(tool_calls_by_turn)
    print("\n  [Skill Adherence]")
    for check, passed in skill_checks.items():
        print(f"    {check}: {'PASS' if passed else 'FAIL'}")

    # Response indicators
    indicator_checks = verify_response_indicators(responses)
    print("\n  [Response Indicators]")
    for check, passed in indicator_checks.items():
        print(f"    {check}: {'PASS' if passed else 'FAIL'}")

    # ── Pass/fail criteria ──
    all_pass = True

    criteria = [
        # Doc parse → entity store
        (db_state["product_count"] >= 1, "At least 1 product extracted"),
        (db_state["model_count"] >= 2, "At least 2 requirement_models extracted"),
        (db_state["item_count"] >= 5, "At least 5 requirement_items extracted"),
        (skill_checks.get("turn1_used_read_docx", False), "Turn 1 used read_docx (skill step 1)"),
        (skill_checks.get("turn1_used_entity_store", False), "Turn 1 used entity_store (skill steps 3-4)"),
        (skill_checks.get("turn2_used_entity_store", False), "Turn 2 used entity_store (multi-turn memory)"),
        # Text2SQL — agent should use query_reqmgmt tool
        (skill_checks.get("turn5_used_query_reqmgmt", False), "Turn 5 used query_reqmgmt (Text2SQL tool)"),
        (skill_checks.get("turn6_used_query_reqmgmt", False), "Turn 6 used query_reqmgmt (Text2SQL complex query)"),
    ]

    print("\n  [Pass/Fail Criteria]")
    for passed, desc in criteria:
        mark = "PASS" if passed else "FAIL"
        print(f"    {mark}: {desc}")
        if not passed:
            all_pass = False

    print(f"\n  {'='*60}")
    if all_pass:
        print("  ALL CHECKS PASSED")
    else:
        print("  SOME CHECKS FAILED — review above")
    print(f"  {'='*60}")

    if not all_pass:
        sys.exit(1)


if __name__ == "__main__":
    asyncio.run(main())
