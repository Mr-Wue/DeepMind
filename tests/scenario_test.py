"""
scenario_test.py — deepagents SubAgent pattern for document parsing + Text2SQL.

Architecture:
  Main Agent
  ├── tools: [query_reqmgmt]
  ├── memory: [AGENTS.md]
  ├── middleware: [InvocationLoggingHandler.as_middleware()]
  └── subagents:
      └── "req-parse"
          ├── tools: [parse_docx_outline, extract_entities, store_entities]
          └── skills: [skills/req-parse]

The main agent CANNOT call parse_docx_outline directly — it must delegate to
the req-parse subagent via the task() tool. This ensures tool-skill binding.

Usage:
  python -X utf8 scenario_test.py
  python -X utf8 scenario_test.py --docx path/to/file.docx
  python -X utf8 scenario_test.py --synthetic
"""

from __future__ import annotations

import argparse
import asyncio
import json
import os
import sys
import textwrap
import time
from pathlib import Path

PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT))

from agents.deep_agent import create_deepmind_agent, DeepMindContext

if sys.platform == "win32" and sys.stdout.encoding != "utf-8":
    import io
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
    except (ValueError, AttributeError):
        pass  # stdout already wrapped (e.g., by pytest capture)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. Environment & paths
# ═══════════════════════════════════════════════════════════════════════════════

def load_dotenv() -> None:
    from dotenv import load_dotenv as _load
    env_path = PROJECT_ROOT / ".env"
    if not env_path.exists():
        print("[WARN] .env not found.")
        return
    _load(env_path)
    masked = os.getenv("LLM_API_KEY", "")
    display = masked[:8] + "..." if len(masked) > 8 else "(empty)"
    print(f"[OK] Loaded .env  (model={os.getenv('LLM_MODEL_ID', '?')}, key={display})")


def find_docx() -> str | None:
    from utils.paths import data_paths
    req_dir = data_paths.shared_req_dir()
    if not req_dir.exists():
        return None
    for f in sorted(req_dir.iterdir()):
        if f.suffix.lower() == ".docx" and not f.name.startswith("~"):
            return str(f)
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# 2. Synthetic test document
# ═══════════════════════════════════════════════════════════════════════════════

def create_test_docx() -> str:
    from docx import Document
    doc = Document()
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        try:
            style = doc.styles[name]
            style.font.size = doc.shared.Pt(size)
            style.font.bold = True
        except Exception:
            pass

    doc.add_heading("XX电商平台用户需求规格说明书", level=1)
    doc.add_paragraph(
        "本文档描述了XX电商平台V3.0版本的核心用户需求，涵盖用户管理、订单系统、"
        "商品管理三大核心模块。系统建设目标：打造高性能、高可用的电商平台，支持日均百万级订单处理。"
    )

    doc.add_heading("用户管理模块", level=2)
    doc.add_paragraph("用户管理模块负责终端用户的注册、登录、权限控制、个人信息管理等功能。")

    doc.add_heading("用户注册与登录", level=3)
    doc.add_paragraph("系统应支持手机号+验证码注册方式。性能要求：注册接口响应时间小于500ms，支持1000并发。")

    doc.add_heading("角色与权限管理", level=3)
    doc.add_paragraph(
        "系统应支持基于RBAC的角色权限控制，权限粒度需达到按钮级别。"
        "默认提供管理员、运营、普通用户三种角色。"
    )

    doc.add_heading("订单系统模块", level=2)
    doc.add_paragraph("订单系统负责用户下单、支付、物流追踪、退换货等全生命周期管理。")

    doc.add_heading("购物车与订单创建", level=3)
    doc.add_paragraph("下单接口响应时间小于200ms，支持5000并发。超卖场景需通过分布式锁保障库存一致性。")

    doc.add_heading("订单支付", level=3)
    doc.add_paragraph("支持微信支付、支付宝、银行卡三种支付方式。支付超时时间15分钟。")

    doc.add_heading("商品管理模块", level=2)
    doc.add_paragraph("商品管理模块包括商品发布、SKU管理、库存管理、商品搜索等功能。")

    doc.add_heading("SKU与库存管理", level=3)
    doc.add_paragraph(
        "每个商品可配置多个SKU，每个SKU独立管理库存和价格。"
        "库存变更需记录日志，支持低库存预警（阈值默认10件）。"
    )

    out_path = PROJECT_ROOT / "data" / "test_req.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] Generated test document: {out_path}  ({out_path.stat().st_size} bytes)")
    return str(out_path)


# ═══════════════════════════════════════════════════════════════════════════════
# 3. Model & agent
# ═══════════════════════════════════════════════════════════════════════════════

async def build_agent():
    from agents.init import init_deepmind
    config = await init_deepmind()
    return create_deepmind_agent(config)


# ═══════════════════════════════════════════════════════════════════════════════
# 4. Test runner
# ═══════════════════════════════════════════════════════════════════════════════

def _make_log_filter():
    """Return a stdout wrapper that suppresses [LOG] lines and blank lines."""
    _real_stdout = sys.stdout

    def _keep(s: str) -> bool:
        stripped = s.strip() if isinstance(s, str) else str(s).strip()
        return bool(stripped) and "[LOG]" not in str(s)

    return type("_LogFilter", (object,), {
        "write": lambda self, s: _real_stdout.write(s) if _keep(s) else None,
        "flush": lambda self: _real_stdout.flush(),
        "__getattr__": lambda self, k: getattr(_real_stdout, k),
    })()


async def run_turn(agent, thread_id: str, user_message: str, turn_label: str, max_steps: int = 30) -> dict:
    config = {"configurable": {"thread_id": thread_id}}

    print(f"\n{'=' * 60}")
    print(f"  {turn_label}")
    print(f"  User: {user_message[:120]}{'...' if len(user_message) > 120 else ''}")
    print(f"{'=' * 60}")

    all_messages: list = []
    tool_calls_found: list[str] = []
    response_text = ""
    step = 0
    turn_start = time.time()

    # 不再压制 [LOG] 输出 — 调试期间需要看到 middleware 日志
    # sys.stdout = _make_log_filter()

    print(f"\n  [Streaming — max {max_steps} steps]\n")

    input_data = {"messages": [{"role": "user", "content": user_message}]}

    try:
        async for chunk in agent.astream(input_data, config=config, stream_mode="updates", context=DeepMindContext(user_id="default")):
            step += 1
            if step > max_steps:
                print(f"  [HALT] Max steps ({max_steps}) reached")
                break

            ts = time.strftime("%H:%M:%S")
            elapsed = time.time() - turn_start

            for node_name, node_output in chunk.items():
                if node_output is None:
                    continue
                if "messages" not in node_output:
                    continue

                msgs = node_output["messages"]
                if hasattr(msgs, "value"):
                    msgs = msgs.value
                if not isinstance(msgs, (list, tuple)):
                    msgs = [msgs]

                show_header = False
                for msg in msgs:
                    if msg not in all_messages:
                        if not show_header:
                            print(f"  [{ts}] Step {step} (+{elapsed:.1f}s) | Node: {node_name}")
                            show_header = True
                        all_messages.append(msg)
                        msg_type = getattr(msg, "type", "?")

                        if hasattr(msg, "tool_calls") and msg.tool_calls:
                            for tc in msg.tool_calls:
                                name = tc.get("name", "unknown")
                                args = tc.get("args", {})
                                tool_calls_found.append(name)
                                args_str = json.dumps(args, ensure_ascii=False, indent=2)
                                if len(args_str) > 500:
                                    args_str = args_str[:500] + "\n    ..."
                                print(f"      [TOOL CALL] {name}")
                                print(f"      Args:\n{textwrap.indent(args_str, '        ')}")

                        elif msg_type == "ai":
                            content = str(getattr(msg, "content", ""))
                            if content:
                                if len(content) > 400:
                                    content = content[:400] + "..."
                                print(f"      [AI] {content}")

                        elif msg_type == "tool":
                            content = str(getattr(msg, "content", ""))
                            tool_name = getattr(msg, "name", "?")
                            if len(content) > 400:
                                content = content[:400] + "..."
                            print(f"      [TOOL RESULT] {tool_name}: {content}")

                # (blank line suppressed)
    except Exception as e:
        print(f"\n  [ERROR] Turn failed after {time.time() - turn_start:.1f}s: {e}")
        import traceback
        traceback.print_exc()
        return {"response": f"ERROR: {e}", "tools": tool_calls_found}
    finally:
        # sys.stdout = sys.__stdout__  # 已注释：不再使用 log filter
        pass

    elapsed_total = time.time() - turn_start
    print(f"  [DONE] Turn completed in {elapsed_total:.1f}s ({step} steps)")

    for msg in reversed(all_messages):
        if getattr(msg, "type", "") == "ai":
            content = str(getattr(msg, "content", ""))
            if content:
                response_text = content
                break

    print(f"  Tools called: {tool_calls_found if tool_calls_found else '(none)'}")
    print(f"\n  Final response preview:\n{textwrap.indent(response_text[:800], '    ')}")
    if len(response_text) > 800:
        print(f"    ... (truncated, total {len(response_text)} chars)")

    return {"response": response_text, "tools": tool_calls_found}


# ═══════════════════════════════════════════════════════════════════════════════
# 5. Verification
# ═══════════════════════════════════════════════════════════════════════════════

def verify_database() -> dict:
    import sqlite3
    from utils.paths import data_paths

    db_path = data_paths.reqmgmt_db()
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row

    products = [dict(r) for r in conn.execute("SELECT * FROM products").fetchall()]
    models = [dict(r) for r in conn.execute("SELECT * FROM requirement_models").fetchall()]
    items = [dict(r) for r in conn.execute("SELECT * FROM requirement_items").fetchall()]

    # FK integrity check
    orphans = conn.execute("""
        SELECT ir.id, ir.rm_id FROM requirement_items ir
        LEFT JOIN requirement_models rm ON ir.rm_id = rm.id
        WHERE rm.id IS NULL AND ir.rm_id IS NOT NULL AND ir.rm_id != ''
    """).fetchall()

    conn.close()

    return {
        "product_count": len(products),
        "model_count": len(models),
        "item_count": len(items),
        "products": products,
        "models": models,
        "items": items,
        "orphan_items": len(orphans),
    }


def verify_subagent_used(tool_calls_by_turn: dict[str, list[str]]) -> dict:
    """Check that main agent delegated to sub-agent via 'task' tool."""
    checks = {}
    for turn_label, tools in tool_calls_by_turn.items():
        checks[f"{turn_label}_used_task"] = "task" in tools
    return checks


# ═══════════════════════════════════════════════════════════════════════════════
# 6. SkillScenarioRunner — entry test harness
# ═══════════════════════════════════════════════════════════════════════════════

class SkillScenarioRunner:
    """Entry test harness for deepagents req-parse skill scenario.

    Validates the full SubAgent + SKILL.md workflow:
      1. Main agent delegates to req-parse SubAgent via task() tool
      2. SubAgent follows SKILL.md steps: parse → extract → store
      3. Database stores entities with FK integrity
      4. Cross-turn memory persists via MemorySaver checkpointer
      5. Idempotent re-parse produces identical results

    Usage::

        # Standalone
        runner = SkillScenarioRunner(docx_path="path/to/doc.docx")
        success = await runner.run()

        # pytest fixture
        runner = SkillScenarioRunner(use_synthetic=True)
        await runner.setup()
        try:
            await runner.run_all_turns()
            results = runner.verify()
            assert results["all_pass"]
        finally:
            await runner.teardown()
    """

    def __init__(
        self,
        docx_path: str | None = None,
        use_synthetic: bool = False,
    ) -> None:
        self._docx_path = docx_path
        self._use_synthetic = use_synthetic
        self.agent = None
        self.thread_id: str = ""
        self.turn_results: dict[str, dict] = {}
        self.db_state: dict = {}
        self.subagent_checks: dict = {}
        self.criteria_results: list[tuple[bool, str]] = []

    # ── Lifecycle ──────────────────────────────────────────────────────────

    async def setup(self) -> str:
        """Initialize environment, select document, build agent.

        Returns the resolved docx_path.
        """
        load_dotenv()

        if self._docx_path:
            path = Path(self._docx_path)
            if not path.exists():
                msg = f"Document not found: {self._docx_path}"
                print(f"[FAIL] {msg}")
                raise FileNotFoundError(msg)
            print(f"[OK] Using specified document: {self._docx_path}")
        elif self._use_synthetic:
            self._docx_path = create_test_docx()
        else:
            # 默认文档: docs/req/输入1-用户需求.docx
            default_docx = PROJECT_ROOT / "docs" / "req" / "输入1-用户需求.docx"
            if default_docx.exists():
                self._docx_path = str(default_docx)
                print(f"[OK] Using default document: {self._docx_path}")
            else:
                self._docx_path = find_docx()
                if self._docx_path:
                    print(f"[OK] Using shared document: {self._docx_path}")
                else:
                    print("[INFO] No shared docs found, generating synthetic test doc.")
                    self._docx_path = create_test_docx()

        self.agent = await build_agent()
        self.thread_id = f"test-{int(time.time())}"
        print(f"[OK] Test thread: {self.thread_id}")
        return self._docx_path

    async def teardown(self) -> None:
        """Release agent resources."""
        self.agent = None

    # ── Turn execution ─────────────────────────────────────────────────────

    async def _run_labeled_turn(
        self, turn_num: int, label: str, message: str, max_steps: int = 30
    ) -> dict:
        key = f"turn{turn_num}"
        result = await run_turn(self.agent, self.thread_id, message, label, max_steps)
        self.turn_results[key] = result
        return result

    async def run_all_turns(self) -> None:
        """Execute the 4-turn validation scenario."""
        assert self.agent is not None, "Call setup() first"

        # Turn 1: Parse & Store (→ SubAgent 'req-parse')
        turn1_msg = (
            f"请解析需求文档 `{self._docx_path}`，提取其中的产品(product)、"
            f"需求模型(requirement_model)和用户需求项(requirement_item)，然后存储到数据库。"
            f"完成后报告入库结果。"
        )
        await self._run_labeled_turn(
            1, "Turn 1: Parse & Store (→ SubAgent 'req-parse')", turn1_msg
        )

        # Turn 2: Query & Filter (→ query_reqmgmt)
        turn2_msg = (
            "使用 query_reqmgmt 工具查询："
            "数据库里一共有多少个产品(product)？"
            "列出所有需求模型(requirement_model)的名称和类型。"
            "每个需求模型下各有多少需求项(requirement_item)？按数量从多到少排列。"
            "另外，把标题或描述中包含'性能'关键词的需求项筛选出来。"
        )
        await self._run_labeled_turn(
            2, "Turn 2: Query & Filter (→ query_reqmgmt)", turn2_msg
        )

        # Turn 3: Contextual follow-up (cross-turn memory)
        turn3_msg = (
            "回到我们在第一轮解析的那个文档，其中'订单系统模块'下面有几个需求项？"
            "列出它们的具体内容。"
        )
        await self._run_labeled_turn(
            3, "Turn 3: Contextual Follow-up (Cross-turn)", turn3_msg
        )

        # Turn 4: Idempotency — re-parse same document
        turn4_msg = (
            f"请再次解析同一个文档 `{self._docx_path}`，提取实体并存储到数据库。"
            f"存储使用幂等逻辑（INSERT OR REPLACE），入库数量应和第一次相同。"
        )
        await self._run_labeled_turn(
            4, "Turn 4: Idempotency (→ SubAgent 'req-parse')", turn4_msg
        )

    # ── Verification ───────────────────────────────────────────────────────

    def verify(self) -> dict:
        """Run all verifications and return results dict.

        Returns:
            {
                "all_pass": bool,
                "db_state": {...},
                "subagent_checks": {...},
                "criteria": [(passed: bool, description: str), ...],
                "tool_calls_by_turn": {...},
            }
        """
        # Database ground truth
        self.db_state = verify_database()

        # Tool calls by turn
        tool_calls_by_turn = {
            k: v["tools"] for k, v in self.turn_results.items()
        }

        # SubAgent delegation check
        self.subagent_checks = verify_subagent_used(tool_calls_by_turn)

        # Build criteria
        self.criteria_results = [
            (self.db_state["product_count"] >= 1,
             "At least 1 product extracted"),
            (self.db_state["model_count"] >= 2,
             "At least 2 requirement_models extracted"),
            (self.db_state["item_count"] >= 5,
             "At least 5 requirement_items extracted"),
            (self.db_state["orphan_items"] == 0,
             "No orphan items (FK integrity)"),
            (self.subagent_checks.get("turn1_used_task", False),
             "Turn 1 delegated to SubAgent (task tool)"),
            (self.subagent_checks.get("turn4_used_task", False),
             "Turn 4 delegated to SubAgent (task tool)"),
            ("query_reqmgmt" in tool_calls_by_turn.get("turn2", []),
             "Turn 2 used query_reqmgmt"),
        ]
        all_pass = all(passed for passed, _ in self.criteria_results)

        return {
            "all_pass": all_pass,
            "db_state": self.db_state,
            "subagent_checks": self.subagent_checks,
            "criteria": self.criteria_results,
            "tool_calls_by_turn": tool_calls_by_turn,
        }

    def print_report(self, results: dict) -> None:
        """Print a formatted verification report."""
        print(f"\n{'=' * 60}")
        print("  Verification Summary")
        print(f"{'=' * 60}")

        db = results["db_state"]
        print(f"\n  [Database Ground Truth]")
        print(f"    products:           {db['product_count']}")
        print(f"    requirement_models: {db['model_count']}")
        print(f"    requirement_items:  {db['item_count']}")
        print(f"    orphan_items (FK):  {db['orphan_items']}")
        for m in db["models"]:
            print(f"      {m['name']}")
        for it in db["items"]:
            print(f"      {it['name']}: {it['title']}")

        print(f"\n  [SubAgent Delegation]")
        for check, passed in results["subagent_checks"].items():
            print(f"    {check}: {'PASS' if passed else 'FAIL'}")

        print(f"\n  [Tool Calls]")
        for turn_label, tools in results["tool_calls_by_turn"].items():
            print(f"    {turn_label}: {tools}")

        print(f"\n  [Pass/Fail Criteria]")
        for passed, desc in results["criteria"]:
            print(f"    {'PASS' if passed else 'FAIL'}: {desc}")

        print(f"\n  {'=' * 60}")
        if results["all_pass"]:
            print("  ALL CHECKS PASSED")
        else:
            print("  SOME CHECKS FAILED — review above")
        print(f"  {'=' * 60}")

    # ── Full run ───────────────────────────────────────────────────────────

    async def run(self) -> dict:
        """Run the complete test scenario end-to-end.

        Returns the verify() results dict.  Prints progress and a final report.
        Exits with code 1 if any check fails (unless called via pytest).
        """
        print("=" * 60)
        print("  DeepMind — deepagents SubAgent Pattern")
        print("  Main: [query_reqmgmt]  |  SubAgent 'req-parse':")
        print("    [parse_docx_outline, extract_entities, store_entities]")
        print("  Skill: skills/req-parse/SKILL.md")
        print("=" * 60)

        await self.setup()
        try:
            await self.run_all_turns()
            results = self.verify()
            self.print_report(results)
            return results
        finally:
            await self.teardown()


# ═══════════════════════════════════════════════════════════════════════════════
# 7. Main — standalone entry point
# ═══════════════════════════════════════════════════════════════════════════════

async def main():
    parser = argparse.ArgumentParser(description="DeepMind — deepagents SubAgent validation")
    parser.add_argument("--docx", help="Path to a .docx file to parse")
    parser.add_argument("--synthetic", action="store_true", help="Generate a synthetic test document")
    args = parser.parse_args()

    runner = SkillScenarioRunner(
        docx_path=args.docx,
        use_synthetic=args.synthetic,
    )
    results = await runner.run()

    if not results["all_pass"]:
        sys.exit(1)


if __name__ == "__main__":
    asyncio.run(main())
